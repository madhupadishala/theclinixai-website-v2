#!/usr/bin/env python3
"""Convert approved ClinixAI PV blog DOCX files into native static HTML articles."""

from __future__ import annotations

import html
import json
import re
import sys
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SITE = ROOT / "website"
OUTPUT = SITE / "insights"
SITE_URL = "https://www.theclinixai.com"
PUBLISHED = "2026-07-23"


@dataclass
class Article:
    source: Path
    title: str
    slug: str
    description: str
    primary_keyword: str
    topic: str
    is_mother: bool
    body: str
    faq: list[tuple[str, str]] = field(default_factory=list)
    word_count: int = 0


def iter_blocks(parent: DocumentType) -> Iterable[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def paragraph_html(paragraph: Paragraph) -> str:
    parts: list[str] = []
    rels = paragraph.part.rels
    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:hyperlink"):
            rid = child.get(qn("r:id"))
            href = rels[rid].target_ref if rid and rid in rels else ""
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            if href:
                parts.append(
                    f'<a href="{html.escape(href, quote=True)}" '
                    'target="_blank" rel="noopener noreferrer">'
                    f"{html.escape(text)}</a>"
                )
            else:
                parts.append(html.escape(text))
        elif child.tag == qn("w:r"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            if not text:
                continue
            escaped = html.escape(text)
            rpr = child.find(qn("w:rPr"))
            if rpr is not None:
                if rpr.find(qn("w:b")) is not None:
                    escaped = f"<strong>{escaped}</strong>"
                if rpr.find(qn("w:i")) is not None:
                    escaped = f"<em>{escaped}</em>"
            parts.append(escaped)
    return "".join(parts).strip()


def plain_paragraph_text(paragraph: Paragraph) -> str:
    return "".join(node.text or "" for node in paragraph._p.iter(qn("w:t"))).strip()


def slugify(value: str) -> str:
    value = value.lower().replace("–", "-").replace("—", "-")
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return re.sub(r"-+", "-", value).strip("-")[:96]


def infer_topic(path: Path) -> str:
    value = str(path).lower()
    mappings = [
        ("topic_3", "ICSR Quality, Medical Review & Submission"),
        ("topic3", "ICSR Quality, Medical Review & Submission"),
        ("t4_", "Signal Management"),
        ("t5_", "Aggregate Safety Reporting"),
        ("t6_", "Risk Management & Benefit–Risk"),
        ("batch7", "PV Quality Management System"),
        ("batch8", "PV Agreements & Partner Governance"),
        ("batch9", "Special Situations"),
        ("batch10", "Clinical Trial Safety"),
        ("batch11", "PV Audits & CAPA"),
        ("batch12", "Pharmacoepidemiology & Real-World Evidence"),
    ]
    for marker, topic in mappings:
        if marker in value:
            return topic
    return "Pharmacovigilance"


def seo_fields(doc: DocumentType) -> dict[str, str]:
    fields: dict[str, str] = {}
    in_package = False
    pending = ""
    for paragraph in doc.paragraphs:
        text = plain_paragraph_text(paragraph)
        style = paragraph.style.name or ""
        if style == "Heading 1" and "seo publishing package" in text.lower():
            in_package = True
            continue
        if not in_package or not text:
            continue
        if style.startswith("Heading 2"):
            pending = text.lower()
            continue
        if pending and style == "Normal":
            fields[pending] = text
            pending = ""
    return fields


def choose_field(fields: dict[str, str], names: tuple[str, ...], fallback: str) -> str:
    for name in names:
        for key, value in fields.items():
            if name in key and value:
                return value
    return fallback


def extract_title(doc: DocumentType) -> str:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            text = plain_paragraph_text(paragraph)
            if text and "seo publishing package" not in text.lower():
                return text
    raise ValueError("No article Heading 1 found")


def table_html(table: Table) -> str:
    rows = []
    for row_index, row in enumerate(table.rows):
        cells = []
        tag = "th" if row_index == 0 else "td"
        for cell in row.cells:
            content = "<br>".join(
                paragraph_html(p) for p in cell.paragraphs if plain_paragraph_text(p)
            )
            cells.append(f"<{tag}>{content}</{tag}>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return '<div class="article-table-wrap"><table>' + "".join(rows) + "</table></div>"


def convert_docx(path: Path) -> Article:
    doc = Document(path)
    title = extract_title(doc)
    fields = seo_fields(doc)
    description = choose_field(
        fields,
        ("meta description",),
        f"An expert TheClinixAI guide to {title.lower()}.",
    )
    keyword = choose_field(
        fields,
        ("primary target keyword", "primary keyword"),
        title.split(":")[0].lower(),
    )
    proposed = choose_field(fields, ("url slug",), "")
    proposed = proposed.rstrip("/").split("/")[-1] if proposed else ""
    slug = slugify(proposed or title.split(":")[0])
    topic = infer_topic(path)
    is_mother = "mother" in path.name.lower()

    blocks: list[str] = []
    faq: list[tuple[str, str]] = []
    current_faq_question: str | None = None
    in_article = False
    in_seo = False
    list_type: str | None = None
    list_items: list[str] = []

    def flush_list() -> None:
        nonlocal list_type, list_items
        if list_type and list_items:
            tag = "ol" if list_type == "number" else "ul"
            blocks.append(f"<{tag}>" + "".join(f"<li>{item}</li>" for item in list_items) + f"</{tag}>")
        list_type = None
        list_items = []

    for block in iter_blocks(doc):
        if isinstance(block, Table):
            if in_article and not in_seo:
                flush_list()
                blocks.append(table_html(block))
            continue

        text = plain_paragraph_text(block)
        style = block.style.name or ""
        if style == "Heading 1" and "seo publishing package" in text.lower():
            flush_list()
            in_seo = True
            continue
        if style == "Heading 1" and text == title:
            in_article = True
            continue
        if not in_article or in_seo or not text:
            continue

        rich = paragraph_html(block)
        if style.startswith("List Bullet"):
            if list_type != "bullet":
                flush_list()
                list_type = "bullet"
            list_items.append(rich)
            continue
        if style.startswith("List Number"):
            if list_type != "number":
                flush_list()
                list_type = "number"
            list_items.append(rich)
            continue

        flush_list()
        if style == "Heading 2":
            blocks.append(f'<h2 id="{slugify(text)}">{rich}</h2>')
            current_faq_question = None
        elif style == "Heading 3":
            blocks.append(f'<h3 id="{slugify(text)}">{rich}</h3>')
            if text.endswith("?"):
                current_faq_question = text
        elif style == "Callout":
            blocks.append(f'<aside class="article-callout">{rich}</aside>')
        else:
            blocks.append(f"<p>{rich}</p>")
            if current_faq_question:
                faq.append((current_faq_question, text))
                current_faq_question = None
    flush_list()

    body = "\n".join(blocks)
    word_count = len(re.findall(r"\b[\w’'-]+\b", re.sub(r"<[^>]+>", " ", body)))
    return Article(path, title, slug, description, keyword, topic, is_mother, body, faq, word_count)


def article_template(article: Article, related: list[Article]) -> str:
    url = f"{SITE_URL}/insights/{article.slug}"
    json_ld: list[dict] = [
        {
            "@context": "https://schema.org",
            "@type": "Article",
            "headline": article.title,
            "description": article.description,
            "datePublished": PUBLISHED,
            "dateModified": PUBLISHED,
            "mainEntityOfPage": url,
            "author": {"@type": "Organization", "name": "TheClinixAI"},
            "publisher": {
                "@type": "Organization",
                "name": "TheClinixAI",
                "url": SITE_URL,
            },
        },
        {
            "@context": "https://schema.org",
            "@type": "BreadcrumbList",
            "itemListElement": [
                {"@type": "ListItem", "position": 1, "name": "Home", "item": SITE_URL},
                {
                    "@type": "ListItem",
                    "position": 2,
                    "name": "Insights",
                    "item": f"{SITE_URL}/insights",
                },
                {"@type": "ListItem", "position": 3, "name": article.title, "item": url},
            ],
        },
    ]
    if article.faq:
        json_ld.append(
            {
                "@context": "https://schema.org",
                "@type": "FAQPage",
                "mainEntity": [
                    {
                        "@type": "Question",
                        "name": question,
                        "acceptedAnswer": {"@type": "Answer", "text": answer},
                    }
                    for question, answer in article.faq
                ],
            }
        )
    related_html = "".join(
        f'<a class="related-article" href="/insights/{item.slug}">'
        f'<span>{html.escape(item.topic)}</span><strong>{html.escape(item.title)}</strong></a>'
        for item in related[:4]
    )
    reading_time = max(4, round(article.word_count / 220))
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>{html.escape(article.title)} | TheClinixAI</title>
  <meta name="description" content="{html.escape(article.description, quote=True)}">
  <meta name="author" content="TheClinixAI">
  <meta name="robots" content="index,follow,max-image-preview:large">
  <link rel="canonical" href="{url}">
  <meta property="og:type" content="article">
  <meta property="og:title" content="{html.escape(article.title, quote=True)}">
  <meta property="og:description" content="{html.escape(article.description, quote=True)}">
  <meta property="og:url" content="{url}">
  <meta property="og:site_name" content="TheClinixAI">
  <meta name="twitter:card" content="summary_large_image">
  <link rel="icon" href="/icon-32.png">
  <link rel="stylesheet" href="/style.css">
  <script type="application/ld+json">{json.dumps(json_ld, ensure_ascii=False)}</script>
</head>
<body class="article-page">
  <a class="skip" href="#article">Skip to article</a>
  <div id="site-header"></div>
  <main id="article">
    <header class="article-hero">
      <div class="article-shell">
        <nav class="article-breadcrumb" aria-label="Breadcrumb">
          <a href="/">Home</a><span>›</span><a href="/insights">Insights</a><span>›</span>
          <span>{html.escape(article.topic)}</span>
        </nav>
        <p class="article-kicker">{html.escape(article.topic)}</p>
        <h1>{html.escape(article.title)}</h1>
        <p class="article-deck">{html.escape(article.description)}</p>
        <div class="article-meta">
          <span>TheClinixAI PV Science &amp; Operations</span>
          <span>{PUBLISHED}</span><span>{reading_time} min read</span>
        </div>
      </div>
    </header>
    <div class="article-layout article-shell">
      <article class="article-content">
        {article.body}
        <aside class="article-disclaimer"><strong>Professional scope:</strong> This article provides scientific and operational pharmacovigilance guidance. Applicable legislation, current regional requirements, company procedures and product-specific obligations must be assessed before regulatory action.</aside>
      </article>
      <aside class="article-sidebar">
        <div class="article-sidebar-card">
          <span>Primary subject</span>
          <strong>{html.escape(article.primary_keyword)}</strong>
          <a href="/contact">Discuss your PV workflow →</a>
        </div>
      </aside>
    </div>
    <section class="related-section">
      <div class="article-shell">
        <p class="article-kicker">CONTINUE THE CLUSTER</p>
        <h2>Related pharmacovigilance guidance</h2>
        <div class="related-grid">{related_html}</div>
      </div>
    </section>
  </main>
  <div id="site-footer"></div>
  <script src="/site.js"></script>
</body>
</html>
"""


def hub_template(articles: list[Article]) -> str:
    groups: dict[str, list[Article]] = {}
    for article in articles:
        groups.setdefault(article.topic, []).append(article)
    sections = []
    tones = ["blue", "cyan", "violet", "teal", "indigo"]
    for cluster_index, (topic, items) in enumerate(groups.items()):
        ordered = sorted(items, key=lambda x: (not x.is_mother, x.title))
        mother = next((item for item in ordered if item.is_mother), ordered[0])
        children = [item for item in ordered if item.slug != mother.slug]
        tone = tones[cluster_index % len(tones)]
        child_cards = "".join(
            f'<a class="cluster-child-card cluster-tone-{tone}" href="/insights/{item.slug}">'
            f'<div class="cluster-card-visual" aria-hidden="true"><span>0{index + 1}</span><i></i><i></i><i></i></div>'
            f'<div class="cluster-child-copy"><span class="cluster-card-label">SPECIALIST GUIDE · {item.word_count:,} WORDS</span>'
            f"<h3>{html.escape(item.title)}</h3><strong>Read specialist guide <b>→</b></strong></div></a>"
            for index, item in enumerate(children[:3])
        )
        sections.append(
            f'<section class="insight-cluster cluster-tone-{tone}"><div class="section-heading">'
            f"<div><p class=\"article-kicker\">TOPIC CLUSTER {cluster_index + 1:02d}</p><h2>{html.escape(topic)}</h2></div>"
            f"<span>{len(items)} articles</span></div>"
            f'<a class="cluster-mother-card" href="/insights/{mother.slug}"><div class="cluster-mother-copy">'
            f'<span class="cluster-card-label">MOTHER GUIDE · {mother.word_count:,} WORDS</span><h3>{html.escape(mother.title)}</h3>'
            f"<p>{html.escape(mother.description)}</p><strong>Enter the complete guide <b>→</b></strong></div>"
            f'<div class="cluster-mother-visual" aria-hidden="true"><span>{cluster_index + 1:02d}</span>'
            f"<div><i></i><i></i><i></i><i></i><i></i></div><small>{html.escape(topic)}</small></div></a>"
            f'<div class="cluster-children">{child_cards}</div></section>'
        )
    return f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Pharmacovigilance Insights & Regulatory Guides | TheClinixAI</title>
<meta name="description" content="Scientific, operational and regulatory pharmacovigilance guidance from TheClinixAI across ICSR processing, signal management, risk management, quality systems, clinical safety and real-world evidence.">
<meta name="robots" content="index,follow,max-image-preview:large"><link rel="canonical" href="{SITE_URL}/insights">
<link rel="icon" href="/icon-32.png"><link rel="stylesheet" href="/style.css"></head>
<body class="insights-index-page"><a class="skip" href="#main">Skip to content</a><div id="site-header"></div>
<main id="main"><header class="insights-index-hero"><div class="shell"><p class="article-kicker">THECLINIXAI PV KNOWLEDGE CENTRE</p>
<h1>Pharmacovigilance decisions must survive scientific and regulatory scrutiny.</h1>
<p>Evidence-led guides written for professionals who must explain not only what was decided—but why.</p>
<div class="article-meta"><span>{len(articles)} published guides</span><span>{len(groups)} topic clusters</span></div></div></header>
<div class="shell insight-clusters">{''.join(sections)}</div></main><div id="site-footer"></div><script src="/site.js"></script></body></html>"""


def discover_docx() -> list[Path]:
    roots = [
        ROOT,
        ROOT / "ClinixAI_Blogs_Topics_4_to_6",
        ROOT / "ClinixAI_Batch7_PV_QMS",
        ROOT / "ClinixAI_Batch8_PV_Agreements",
        ROOT / "ClinixAI_Batch9_Special_Situations",
        ROOT / "ClinixAI_Batch10_Clinical_Trial_Safety",
        ROOT / "ClinixAI_Batch11_PV_Audits_CAPA",
        ROOT / "ClinixAI_Batch12_Pharmacoepidemiology_RWE",
    ]
    paths: list[Path] = []
    for root in roots:
        if not root.exists():
            continue
        for path in root.glob("*.docx"):
            name = path.name.lower()
            if "publishing_index" in name or name.startswith("00_"):
                continue
            if (
                "topic3" in name
                or "topic_3" in name
                or re.match(r"t[456]_", name)
                or re.match(r"b(?:7|8|9|10|11|12)_", name)
            ):
                paths.append(path)
    return sorted(set(paths))


def update_sitemap(articles: list[Article]) -> None:
    existing = [
        "",
        "nexus-platform",
        "services",
        "academy",
        "resources",
        "insights",
        "about",
        "contact",
        "research-001-beyond-automation",
        "research-002-compliant-ai-literature-screening",
        "research-003-ai-nexus-intake-engine",
    ]
    urls = existing + [f"insights/{article.slug}" for article in articles]
    entries = "".join(
        f"<url><loc>{SITE_URL}/{html.escape(path)}</loc><lastmod>{PUBLISHED}</lastmod></url>"
        for path in urls
    )
    (SITE / "sitemap.xml").write_text(
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
        + entries
        + "</urlset>",
        encoding="utf-8",
    )


def main() -> int:
    paths = discover_docx()
    articles = [convert_docx(path) for path in paths]
    slugs = [article.slug for article in articles]
    if len(slugs) != len(set(slugs)):
        raise RuntimeError("Duplicate article slugs detected")
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for article in articles:
        related = [
            candidate
            for candidate in articles
            if candidate.slug != article.slug and candidate.topic == article.topic
        ]
        if len(related) < 4:
            related += [
                candidate
                for candidate in articles
                if candidate.slug != article.slug and candidate not in related
            ][: 4 - len(related)]
        (OUTPUT / f"{article.slug}.html").write_text(
            article_template(article, related), encoding="utf-8"
        )
    (OUTPUT / "index.html").write_text(hub_template(articles), encoding="utf-8")
    update_sitemap(articles)
    manifest = [
        {
            "title": item.title,
            "slug": item.slug,
            "topic": item.topic,
            "mother": item.is_mother,
            "words": item.word_count,
            "source": str(item.source.relative_to(ROOT)),
        }
        for item in articles
    ]
    (OUTPUT / "articles.json").write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    print(f"Generated {len(articles)} articles across {len(set(a.topic for a in articles))} clusters")
    print(f"Total article words: {sum(a.word_count for a in articles):,}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
